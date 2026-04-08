from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def _safe_str(value):
    if value is None:
        return ""
    return str(value)


def _replace_token_in_paragraph(paragraph, token_map: dict):
    if not paragraph.runs:
        return

    full_text = "".join(run.text for run in paragraph.runs)
    new_text = full_text

    for key, value in token_map.items():
        token = f"{{{{{key}}}}}"
        new_text = new_text.replace(token, _safe_str(value))

    if new_text == full_text:
        return

    run_lengths = [len(run.text) for run in paragraph.runs]

    pos = 0
    for i, run in enumerate(paragraph.runs):
        length = run_lengths[i]
        run.text = new_text[pos:pos + length]
        pos += length

    if pos < len(new_text):
        paragraph.runs[-1].text += new_text[pos:]


def _replace_token_in_table(table, token_map: dict):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_token_in_paragraph(paragraph, token_map)


def _set_paragraph_format(paragraph, *, align=None, font_size=None, bold=None, reset_indent=False):
    if align is not None:
        paragraph.alignment = align

    if reset_indent:
        fmt = paragraph.paragraph_format
        fmt.left_indent = Pt(0)
        fmt.first_line_indent = Pt(0)
        fmt.right_indent = Pt(0)

    if not paragraph.runs:
        paragraph.add_run("")

    for run in paragraph.runs:
        if font_size is not None:
            run.font.size = Pt(font_size)
        if bold is not None:
            run.font.bold = bold

def _remove_paragraph(paragraph):
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def _cleanup_trailing_empty_paragraphs(doc):
    for paragraph in reversed(doc.paragraphs):
        if paragraph.text.strip():
            break
        _remove_paragraph(paragraph)


def _normalize_text(text: str) -> str:
    return " ".join(text.replace("\xa0", " ").split()).strip()


def _apply_fixed_format(doc: Document):
    for paragraph in doc.paragraphs:
        raw = paragraph.text.strip()
        normalized = _normalize_text(paragraph.text)

        # 제목
        if normalized == "이 행 확 인 서":
            _set_paragraph_format(
                paragraph,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            continue

        # 가/나/다/라 항목
        if re.match(r"^[가나다라]\.", raw):
            _set_paragraph_format(
                paragraph,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                font_size=14,
                bold=False,
                reset_indent=True,
            )
            continue

        # 비고 첫 줄
        if raw.startswith("※"):
            _set_paragraph_format(
                paragraph,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                font_size=11,
                bold=False,
                reset_indent=True,
            )
            continue


def generate_doc_10110_a(template_path: Path, output_path: Path, project_data: dict):
    doc = Document(template_path)
    token_map = project_data.get("fields", {})

    for paragraph in doc.paragraphs:
        _replace_token_in_paragraph(paragraph, token_map)

    for table in doc.tables:
        _replace_token_in_table(table, token_map)

    _apply_fixed_format(doc)
    _cleanup_trailing_empty_paragraphs(doc)
    doc.save(output_path)