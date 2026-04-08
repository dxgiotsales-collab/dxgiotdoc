from pathlib import Path
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from engines.token_engine import replace_token_in_paragraph, replace_token_in_table
from PIL import Image, ImageOps


def _normalize_image_orientation(image_path: str) -> str:
    img = Image.open(image_path)
    img = ImageOps.exif_transpose(img)

    stem = Path(image_path).stem
    suffix = Path(image_path).suffix.lower()
    fixed_path = Path("app_data/uploads") / f"{stem}_fixed{suffix}"

    img.save(fixed_path)
    return str(fixed_path)


def _get_fit_size_cm(image_path: str, max_width_cm: float = 12, max_height_cm: float = 12):
    with Image.open(image_path) as img:
        width_px, height_px = img.size

    if width_px == 0 or height_px == 0:
        return max_width_cm, max_height_cm

    width_cm = width_px / 96 * 2.54
    height_cm = height_px / 96 * 2.54

    scale = min(max_width_cm / width_cm, max_height_cm / height_cm, 1.0)

    return width_cm * scale, height_cm * scale


def insert_image_to_cell(
    cell,
    image_path: str,
    keep_text: str | None = None,
    max_width_cm: float = 12,
    max_height_cm: float = 12,
):
    if not image_path:
        return

    fixed_image_path = _normalize_image_orientation(image_path)
    fit_width_cm, fit_height_cm = _get_fit_size_cm(
        fixed_image_path,
        max_width_cm=max_width_cm,
        max_height_cm=max_height_cm,
    )

    cell.text = ""

    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if keep_text:
        p1.add_run(keep_text)
        cell.add_paragraph("")

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run()
    run2.add_picture(
        fixed_image_path,
        width=Cm(fit_width_cm),
        height=Cm(fit_height_cm),
    )


def generate_image_doc(template_path: Path, output_path: Path, project_data: dict, doc_code: str):
    from pathlib import Path

    doc = Document(template_path)
    token_map = project_data.get("fields", {})

    for paragraph in doc.paragraphs:
        replace_token_in_paragraph(paragraph, token_map)

    for table in doc.tables:
        replace_token_in_table(table, token_map)

    if doc_code == "DOC_10050":
        image_value = project_data.get("images", {}).get("BUSINESS_LOCATION_MAP_FILE")

        print("DOC_10050 raw image value =", image_value)

        image_path = None

        if isinstance(image_value, str) and image_value.strip():
            image_path = image_value.strip()

        elif hasattr(image_value, "getbuffer") and hasattr(image_value, "name"):
            save_dir = Path("app_data/uploads/runtime")
            save_dir.mkdir(parents=True, exist_ok=True)
            save_path = save_dir / image_value.name
            with open(save_path, "wb") as f:
                f.write(image_value.getbuffer())
            image_path = str(save_path)

        if image_path:
            if Path(image_path).exists():
                pass
            else:
                candidate_1 = Path("app_data/uploads") / image_path
                if candidate_1.exists():
                    image_path = str(candidate_1)
                else:
                    candidate_2 = Path("app_data/uploads/attachments") / Path(image_path).name
                    if candidate_2.exists():
                        image_path = str(candidate_2)
                    else:
                        candidate_3 = Path("app_data/uploads/runtime_10024") / Path(image_path).name
                        if candidate_3.exists():
                            image_path = str(candidate_3)

        print("DOC_10050 resolved image_path =", image_path)
        print("DOC_10050 image exists =", Path(image_path).exists() if image_path else False)

        if image_path and Path(image_path).exists() and len(doc.tables) >= 1:
            table = doc.tables[0]
            last_row = table.rows[-1]
            merged_cell = last_row.cells[0]
            for i in range(1, len(last_row.cells)):
                merged_cell = merged_cell.merge(last_row.cells[i])

            insert_image_to_cell(
                merged_cell,
                image_path,
                keep_text="<약   도> (※ 인근 주거지역과의 이격거리 반드시 표시-카카오맵 지도 활용)",
                max_width_cm=17,
                max_height_cm=12,
            )
        else:
            print("DOC_10050 image insertion skipped")

    elif doc_code == "DOC_10022":
        image_value = project_data.get("images", {}).get("INSTALL_LAYOUT_FILE")

        print("DOC_10022 raw image value =", image_value)
        print("DOC_10022 raw image type =", type(image_value))
        print("DOC_10022 table count =", len(doc.tables))

        image_path = None

        if isinstance(image_value, str) and image_value.strip():
            image_path = image_value.strip()

        elif hasattr(image_value, "getbuffer") and hasattr(image_value, "name"):
            save_dir = Path("app_data/uploads/runtime")
            save_dir.mkdir(parents=True, exist_ok=True)
            save_path = save_dir / image_value.name
            with open(save_path, "wb") as f:
                f.write(image_value.getbuffer())
            image_path = str(save_path)

        if image_path and not Path(image_path).exists():
            candidate_1 = Path("app_data/uploads") / image_path
            if candidate_1.exists():
                image_path = str(candidate_1)
            else:
                candidate_2 = Path("app_data/uploads/attachments") / Path(image_path).name
                if candidate_2.exists():
                    image_path = str(candidate_2)
                else:
                    candidate_3 = Path("app_data/uploads/runtime_10024") / Path(image_path).name
                    if candidate_3.exists():
                        image_path = str(candidate_3)

        print("DOC_10022 resolved image_path =", image_path)
        print("DOC_10022 image exists =", Path(image_path).exists() if image_path else False)

        if image_path and Path(image_path).exists() and len(doc.tables) >= 1:
            table = doc.tables[-1]
            last_row = table.rows[-1]
            merged_cell = last_row.cells[0]
            for i in range(1, len(last_row.cells)):
                merged_cell = merged_cell.merge(last_row.cells[i])

            insert_image_to_cell(
                merged_cell,
                image_path,
                max_width_cm=17,
                max_height_cm=12,
            )
        else:
            print("DOC_10022 image insertion skipped")

    doc.save(output_path)