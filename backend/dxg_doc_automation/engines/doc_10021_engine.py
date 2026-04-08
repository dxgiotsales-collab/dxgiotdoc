from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.table import _Row


def _safe_str(value):
    if value is None:
        return ""
    return str(value)


def _set_cell_align_center(cell):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _replace_text_in_paragraph(paragraph, replacements: dict):
    full_text = "".join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text

    for key, value in replacements.items():
        token = f"{{{{{key}}}}}"
        if token in full_text:
            full_text = full_text.replace(token, _safe_str(value))

    if paragraph.runs:
        paragraph.runs[0].text = full_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = full_text


def _replace_text_in_cell(cell, replacements: dict):
    for paragraph in cell.paragraphs:
        _replace_text_in_paragraph(paragraph, replacements)


def _replace_text_in_table(table, replacements: dict):
    for row in table.rows:
        for cell in row.cells:
            _replace_text_in_cell(cell, replacements)


def _fill_sensor_basis_table(doc: Document, sensor_basis_items: list[dict]):
    target_table = None
    template_row = None

    for table in doc.tables:
        for row in table.rows:
            row_text = "".join(cell.text for cell in row.cells)
            if "{{ITEM_NAME}}" in row_text and "{{BASIS_TEXT}}" in row_text:
                target_table = table
                template_row = row
                break
        if target_table is not None:
            break

    if target_table is None or template_row is None:
        return

    table = target_table
    tbl = table._tbl

    template_tr = deepcopy(template_row._tr)
    insert_idx = list(tbl).index(template_row._tr)

    tbl.remove(template_row._tr)

    for item in sensor_basis_items:
        new_tr = deepcopy(template_tr)
        tbl.insert(insert_idx, new_tr)

        new_row = _Row(new_tr, table)

        replacements = {
            "ITEM_NAME": item.get("ITEM_NAME", ""),
            "BASIS_TEXT": "",   # ← 무조건 공란
        }

        for cell in new_row.cells:
            full_text = cell.text
            for key, value in replacements.items():
                full_text = full_text.replace(f"{{{{{key}}}}}", _safe_str(value))
            cell.text = full_text
            _set_cell_align_center(cell)

        insert_idx += 1


def generate_doc_10021(template_path: Path, output_path: Path, project_data: dict):
    doc = Document(template_path)

    field_map = project_data.get("fields", {})
    sensor_basis_items = project_data.get("sensor_basis_items", [])

    for paragraph in doc.paragraphs:
        _replace_text_in_paragraph(paragraph, field_map)

    for table in doc.tables:
        _replace_text_in_table(table, field_map)

    _fill_sensor_basis_table(doc, sensor_basis_items)

    for paragraph in doc.paragraphs:
        _replace_text_in_paragraph(paragraph, field_map)

    for table in doc.tables:
        _replace_text_in_table(table, field_map)

    doc.save(output_path)