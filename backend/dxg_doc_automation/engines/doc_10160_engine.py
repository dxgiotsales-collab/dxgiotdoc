from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _safe_str(value):
    if value is None:
        return ""
    return str(value)


def _replace_token_in_paragraph(paragraph, token_map: dict):
    if not paragraph.runs:
        return

    full_text = "".join(run.text for run in paragraph.runs)
    new_text = full_text

    for key, value in token_map.items():
        token = f"{{{{{key}}}}}"
        new_text = new_text.replace(token, _safe_str(value))

    if new_text == full_text:
        return

    run_lengths = [len(run.text) for run in paragraph.runs]

    pos = 0
    for i, run in enumerate(paragraph.runs):
        length = run_lengths[i]
        run.text = new_text[pos:pos + length]
        pos += length

    if pos < len(new_text):
        paragraph.runs[-1].text += new_text[pos:]


def _replace_token_in_table(table, token_map: dict):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_token_in_paragraph(paragraph, token_map)


def _remove_paragraph(paragraph):
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def _cleanup_trailing_empty_paragraphs(doc):
    for paragraph in reversed(doc.paragraphs):
        text = paragraph.text.replace("\xa0", "").strip()
        if text:
            break
        _remove_paragraph(paragraph)


def _normalize_text(text: str) -> str:
    return "".join(text.split())


def _fix_title_alignment(doc: Document):
    for paragraph in doc.paragraphs:
        if _normalize_text(paragraph.text) == "위임장":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            break


def generate_doc_10160(template_path: Path, output_path: Path, project_data: dict):
    doc = Document(template_path)
    token_map = project_data.get("fields", {})

    for paragraph in doc.paragraphs:
        _replace_token_in_paragraph(paragraph, token_map)

    for table in doc.tables:
        _replace_token_in_table(table, token_map)

    _fix_title_alignment(doc)
    _cleanup_trailing_empty_paragraphs(doc)
    doc.save(output_path)