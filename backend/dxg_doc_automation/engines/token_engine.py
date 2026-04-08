from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def _safe_str(value):
    if value is None:
        return ""
    return str(value)

def fix_known_title_alignment(doc):
    target_titles = {
        "이 행 확 인 서",
        "위       임       장",
        "사후관리 이행 동의서",
        "보조금 반납 확약서",
    }

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in target_titles:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def replace_token_in_paragraph(paragraph, token_map: dict):
    if not paragraph.runs:
        return

    full_text = "".join(run.text for run in paragraph.runs)
    new_text = full_text

    for key, value in token_map.items():
        token = f"{{{{{key}}}}}"
        new_text = new_text.replace(token, _safe_str(value))

    # 치환할 게 없으면 그대로 종료
    if new_text == full_text:
        return

    # 원본 run 길이 정보
    run_lengths = [len(run.text) for run in paragraph.runs]

    # 새 텍스트를 기존 run 길이대로 다시 분배
    pos = 0
    for i, run in enumerate(paragraph.runs):
        length = run_lengths[i]
        run.text = new_text[pos:pos + length]
        pos += length

    # 남는 텍스트는 마지막 run에 붙임
    if pos < len(new_text):
        paragraph.runs[-1].text += new_text[pos:]


def replace_token_in_table(table, token_map: dict):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_token_in_paragraph(paragraph, token_map)


def remove_paragraph(paragraph):
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def cleanup_trailing_empty_paragraphs(doc):
    for paragraph in reversed(doc.paragraphs):
        if paragraph.text.strip():
            break
        remove_paragraph(paragraph)


def generate_token_doc(template_path: Path, output_path: Path, project_data: dict):
    doc = Document(template_path)

    token_map = project_data.get("fields", {})

    for paragraph in doc.paragraphs:
        replace_token_in_paragraph(paragraph, token_map)

    for table in doc.tables:
        replace_token_in_table(table, token_map)

    fix_known_title_alignment(doc)
    cleanup_trailing_empty_paragraphs(doc)

    doc.save(output_path)