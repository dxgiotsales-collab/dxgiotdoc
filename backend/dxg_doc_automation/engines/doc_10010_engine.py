from pathlib import Path
from docx import Document


def _safe_str(value):
    if value is None:
        return ""
    return str(value)


def _replace_text_in_paragraph(paragraph, replacements: dict):
    full_text = "".join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text

    for key, value in replacements.items():
        token = f"{{{{{key}}}}}"
        if token in full_text:
            full_text = full_text.replace(token, _safe_str(value))

    if paragraph.runs:
        paragraph.runs[0].text = full_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = full_text


def _replace_text_in_cell(cell, replacements: dict):
    for paragraph in cell.paragraphs:
        _replace_text_in_paragraph(paragraph, replacements)


def _replace_text_in_table(table, replacements: dict):
    for row in table.rows:
        for cell in row.cells:
            _replace_text_in_cell(cell, replacements)


def _find_template_row(table, token_keys: list[str]):
    for row in table.rows:
        row_text = "".join(cell.text for cell in row.cells)
        if any(f"{{{{{key}}}}}" in row_text for key in token_keys):
            return row
    return None


def _fill_repeat_table(table, items: list[dict], token_keys: list[str]):
    template_row = _find_template_row(table, token_keys)
    if template_row is None:
        return

    template_texts = [cell.text for cell in template_row.cells]
    tbl = table._tbl
    tbl.remove(template_row._tr)

    for item in items:
        new_row = table.add_row()

        for i, cell in enumerate(new_row.cells):
            if i >= len(template_texts):
                continue

            rendered = template_texts[i]
            for key in token_keys:
                rendered = rendered.replace(f"{{{{{key}}}}}", _safe_str(item.get(key, "")))

            cell.text = rendered


def generate_doc_10010(template_path: Path, output_path: Path, project_data: dict, doc_code: str):
    doc = Document(template_path)

    field_map = project_data.get("fields", {})

    # 1) 공통 FIELD 치환
    for paragraph in doc.paragraphs:
        _replace_text_in_paragraph(paragraph, field_map)

    for table in doc.tables:
        _replace_text_in_table(table, field_map)

    # 2) A안
    if doc_code == "DOC_10010_A":
        site_facility_status = project_data.get("site_facility_status", [])
        site_facility_tokens = [
            "EMISSION_FACILITY_NAME",
            "EMISSION_CAPACITY",
            "EMISSION_QTY",
            "PREVENTION_METHOD",
            "PREVENTION_CAPACITY",
            "PREVENTION_QTY",
        ]

        install_items = project_data.get("install_items", [])
        install_item_tokens = [
            "ITEM_NAME",
            "ITEM_UNIT_PRICE",
            "ITEM_QTY",
            "ITEM_AMOUNT",
        ]

        for table in doc.tables:
            _fill_repeat_table(table, site_facility_status, site_facility_tokens)

        for table in doc.tables:
            _fill_repeat_table(table, install_items, install_item_tokens)

    # 3) B안
    elif doc_code == "DOC_10010_B":
        pollutants = project_data.get("pollutants", [])
        pollutant_tokens = [
            "ITEM_POLLUTANT_TYPE",
            "ITEM_POLLUTANT_AMOUNT",
        ]

        for table in doc.tables:
            _fill_repeat_table(table, pollutants, pollutant_tokens)

    # 4) 마지막 한 번 더 FIELD 치환
    for paragraph in doc.paragraphs:
        _replace_text_in_paragraph(paragraph, field_map)

    for table in doc.tables:
        _replace_text_in_table(table, field_map)

    doc.save(output_path)