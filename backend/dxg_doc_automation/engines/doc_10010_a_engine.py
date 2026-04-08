from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.table import _Row


def _safe_str(value):
    if value is None:
        return ""
    if isinstance(value, int):
        return f"{value:,}"
    return str(value)


def _normalize_text(text: str) -> str:
    return str(text or "").replace("\n", "").replace(" ", "").strip()


def _set_cell_align_center(cell):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _replace_text_in_paragraph(paragraph, replacements: dict):
    full_text = "".join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text

    for key, value in replacements.items():
        token = f"{{{{{key}}}}}"
        if token in full_text:
            full_text = full_text.replace(token, _safe_str(value))

    if paragraph.runs:
        paragraph.runs[0].text = full_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = full_text


def _replace_text_in_cell(cell, replacements: dict):
    for paragraph in cell.paragraphs:
        _replace_text_in_paragraph(paragraph, replacements)


def _replace_text_in_table(table, replacements: dict):
    for row in table.rows:
        for cell in row.cells:
            _replace_text_in_cell(cell, replacements)


def _find_site_facility_target(doc: Document):
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            row_text = "".join(cell.text for cell in row.cells)
            if "{{BLOCK_SITE_FACILITY_STATUS}}" not in row_text:
                continue

            header_row = None
            for up_idx in range(row_idx - 1, -1, -1):
                candidate = table.rows[up_idx]
                candidate_text = _normalize_text("".join(cell.text for cell in candidate.cells))
                if (
                    "배출시설명" in candidate_text
                    or "방지시설종류" in candidate_text
                    or "최근자가측정결과" in candidate_text
                    or "자가측정결과" in candidate_text
                ):
                    header_row = candidate
                    break

            if header_row is None and row_idx - 1 >= 0:
                header_row = table.rows[row_idx - 1]

            return table, row, header_row

    return None, None, None


def _find_col_idx_from_header(header_texts, keywords):
    if isinstance(keywords, str):
        keywords = [keywords]

    normalized_keywords = [_normalize_text(k) for k in keywords]

    for i, txt in enumerate(header_texts):
        for keyword in normalized_keywords:
            if keyword and keyword in txt:
                return i
    return None

def _fill_install_items_table(doc: Document, install_items: list[dict]):
    target_table = None
    template_row = None
    total_row = None

    for table in doc.tables:
        temp_template_row = None
        temp_total_row = None

        for row in table.rows:
            row_text = "".join(cell.text for cell in row.cells)

            if "{{ITEM_NAME}}" in row_text and "{{ITEM_UNIT_PRICE}}" in row_text:
                temp_template_row = row

            if "합" in row_text and "{{TOTAL_COST}}" in row_text:
                temp_total_row = row

        if temp_template_row is not None and temp_total_row is not None:
            target_table = table
            template_row = temp_template_row
            total_row = temp_total_row
            break

    if target_table is None or template_row is None or total_row is None:
        return

    table = target_table
    tbl = table._tbl

    template_tr = deepcopy(template_row._tr)
    total_tr = deepcopy(total_row._tr)

    template_idx = list(tbl).index(template_row._tr)
    total_cost = sum(int(item.get("ITEM_AMOUNT", 0) or 0) for item in install_items)

    tbl.remove(template_row._tr)
    tbl.remove(total_row._tr)

    insert_idx = template_idx
    created_rows = []

    for item in install_items:
        new_tr = deepcopy(template_tr)
        tbl.insert(insert_idx, new_tr)

        new_row = _Row(new_tr, table)
        created_rows.append(new_row)

        replacements = {
            "ITEM_NAME": item.get("ITEM_NAME", ""),
            "ITEM_UNIT_PRICE": item.get("ITEM_UNIT_PRICE", ""),
            "ITEM_QTY": item.get("ITEM_QTY", ""),
            "ITEM_AMOUNT": item.get("ITEM_AMOUNT", ""),
        }

        for cell in new_row.cells:
            full_text = cell.text
            for key, value in replacements.items():
                full_text = full_text.replace(f"{{{{{key}}}}}", _safe_str(value))
            cell.text = full_text
            _set_cell_align_center(cell)

        insert_idx += 1

    new_total_tr = deepcopy(total_tr)
    tbl.insert(insert_idx, new_total_tr)

    new_total_row = _Row(new_total_tr, table)

    for cell in new_total_row.cells:
        full_text = cell.text.replace("{{TOTAL_COST}}", _safe_str(total_cost))
        cell.text = full_text
        _set_cell_align_center(cell)

    if install_items:
        remark_col_idx = len(new_total_row.cells) - 1
        first_remark_cell = created_rows[0].cells[remark_col_idx]
        merged_remark_cell = first_remark_cell

        for row in created_rows[1:]:
            merged_remark_cell = merged_remark_cell.merge(row.cells[remark_col_idx])

        merged_remark_cell = merged_remark_cell.merge(new_total_row.cells[remark_col_idx])
        merged_remark_cell.text = "단위 : 원\n(VAT 제외)"
        _set_cell_align_center(merged_remark_cell)

def _format_measurement_item_line(item: dict) -> str:
    pollutant = _safe_str(item.get("pollutant", ""))
    amount = _safe_str(item.get("amount", ""))
    unit = _safe_str(item.get("unit", ""))
    date = _safe_str(item.get("date", ""))

    line1 = f"{pollutant} {amount}{unit}".strip()

    if date:
        try:
            yyyy, mm, dd = date.split("-")
            line2 = f"({yyyy[2:]}.{mm}.{dd})"
        except Exception:
            line2 = f"({date})"
        return f"{line1}\n{line2}".strip()

    return line1


def _fill_site_facility_status_table(
    doc: Document,
    site_facility_status: list[dict],
    measurement_items: list[dict],
):
    table, template_row, header_row = _find_site_facility_target(doc)

    if table is None or template_row is None:
        return

    tbl = table._tbl
    template_tr = deepcopy(template_row._tr)
    insert_idx = list(tbl).index(template_row._tr)

    tbl.remove(template_row._tr)

    if not site_facility_status:
        return

    header_texts = []
    if header_row is not None:
        header_texts = [_normalize_text(c.text) for c in header_row.cells]

    prevention_method_col_idx = _find_col_idx_from_header(header_texts, ["방지시설종류", "방지시설"])
    measure_col_idx = _find_col_idx_from_header(header_texts, ["최근자가측정결과", "자가측정결과"])

    qty_indices = [i for i, txt in enumerate(header_texts) if "수량" in txt]
    cap_indices = [i for i, txt in enumerate(header_texts) if "용량" in txt]

    prevention_capacity_col_idx = cap_indices[-1] if len(cap_indices) >= 2 else None
    prevention_qty_col_idx = qty_indices[-1] if len(qty_indices) >= 2 else None

    created_rows = []

    for idx, item in enumerate(site_facility_status):
        new_tr = deepcopy(template_tr)
        tbl.insert(insert_idx, new_tr)

        new_row = _Row(new_tr, table)
        created_rows.append(new_row)

        replacements = {
            "BLOCK_SITE_FACILITY_STATUS": "",
            "EMISSION_FACILITY_NAME": item.get("EMISSION_FACILITY_NAME", ""),
            "EMISSION_CAPACITY": item.get("EMISSION_CAPACITY", ""),
            "EMISSION_QTY": item.get("EMISSION_QTY", ""),
            "PREVENTION_METHOD": item.get("PREVENTION_METHOD", item.get("PREVENTION_FACILITY_NAME", "")),
            "PREVENTION_CAPACITY": item.get("PREVENTION_CAPACITY", ""),
            "PREVENTION_QTY": item.get("PREVENTION_QTY", ""),
            "MEASURE_DATE": "",
        }

        for cell in new_row.cells:
            full_text = cell.text
            for key, value in replacements.items():
                full_text = full_text.replace(f"{{{{{key}}}}}", _safe_str(value))
            cell.text = full_text
            _set_cell_align_center(cell)

        insert_idx += 1

    if not created_rows:
        return

    # 자가측정결과 열 병합 (measurement_items 기준)
    if measure_col_idx is not None and created_rows and measurement_items:
        total_rows = len(created_rows)
        total_items = len(measurement_items)

        base_span = total_rows // total_items
        extra = total_rows % total_items

        row_cursor = 0

        for item_idx, item in enumerate(measurement_items):
            span = base_span + (1 if item_idx < extra else 0)
            start_row = row_cursor
            end_row = min(total_rows - 1, row_cursor + span - 1)

            if start_row > end_row or start_row >= total_rows:
                continue

            merged_cell = created_rows[start_row].cells[measure_col_idx]

            for row_idx in range(start_row + 1, end_row + 1):
                created_rows[row_idx].cells[measure_col_idx].text = ""
                merged_cell = merged_cell.merge(created_rows[row_idx].cells[measure_col_idx])

            merged_cell.text = _format_measurement_item_line(item)
            _set_cell_align_center(merged_cell)

            row_cursor = end_row + 1

    # 방지시설 종류 / 용량 / 수량 병합
    if (
        prevention_method_col_idx is not None
        and prevention_capacity_col_idx is not None
        and prevention_qty_col_idx is not None
    ):
        merge_start = 0
        current_key = (
            site_facility_status[0].get("PREVENTION_METHOD", ""),
            site_facility_status[0].get("PREVENTION_CAPACITY", ""),
            site_facility_status[0].get("PREVENTION_QTY", ""),
        )

        for i in range(1, len(site_facility_status) + 1):
            if i < len(site_facility_status):
                next_key = (
                    site_facility_status[i].get("PREVENTION_METHOD", ""),
                    site_facility_status[i].get("PREVENTION_CAPACITY", ""),
                    site_facility_status[i].get("PREVENTION_QTY", ""),
                )
            else:
                next_key = None

            if next_key != current_key:
                if i - merge_start > 1:
                    merged_cell = created_rows[merge_start].cells[prevention_method_col_idx]
                    for row_idx in range(merge_start + 1, i):
                        created_rows[row_idx].cells[prevention_method_col_idx].text = ""
                        merged_cell = merged_cell.merge(created_rows[row_idx].cells[prevention_method_col_idx])
                    merged_cell.text = _safe_str(current_key[0])
                    _set_cell_align_center(merged_cell)

                    merged_cell = created_rows[merge_start].cells[prevention_capacity_col_idx]
                    for row_idx in range(merge_start + 1, i):
                        created_rows[row_idx].cells[prevention_capacity_col_idx].text = ""
                        merged_cell = merged_cell.merge(created_rows[row_idx].cells[prevention_capacity_col_idx])
                    merged_cell.text = _safe_str(current_key[1])
                    _set_cell_align_center(merged_cell)

                    merged_cell = created_rows[merge_start].cells[prevention_qty_col_idx]
                    for row_idx in range(merge_start + 1, i):
                        created_rows[row_idx].cells[prevention_qty_col_idx].text = ""
                        merged_cell = merged_cell.merge(created_rows[row_idx].cells[prevention_qty_col_idx])
                    merged_cell.text = _safe_str(current_key[2])
                    _set_cell_align_center(merged_cell)

                if i < len(site_facility_status):
                    merge_start = i
                    current_key = next_key


def generate_doc_10010_a(template_path: Path, output_path: Path, project_data: dict):
    doc = Document(template_path)

    field_map = project_data.get("fields", {})

    field_map_no_total = dict(field_map)
    field_map_no_total.pop("TOTAL_COST", None)
    field_map_no_total.pop("MEASURE_DATE", None)

    site_facility_status = project_data.get("site_facility_status", [])
    install_items = project_data.get("install_items", [])

    business = project_data.get("business", {}) or {}
    measurement_items = business.get("measurementItems", []) or []

    for paragraph in doc.paragraphs:
        _replace_text_in_paragraph(paragraph, field_map_no_total)

    for table in doc.tables:
        _replace_text_in_table(table, field_map_no_total)

    _fill_site_facility_status_table(doc, site_facility_status, measurement_items)
    _fill_install_items_table(doc, install_items)

    final_field_map = dict(field_map)
    final_field_map.pop("MEASURE_DATE", None)

    for paragraph in doc.paragraphs:
        _replace_text_in_paragraph(paragraph, final_field_map)

    for table in doc.tables:
        _replace_text_in_table(table, final_field_map)

    doc.save(output_path)