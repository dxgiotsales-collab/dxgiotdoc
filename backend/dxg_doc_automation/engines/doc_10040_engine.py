from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL


def _set_cell_bold(cell, bold=True, font_size=12):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = bold
            run.font.size = Pt(font_size)


def _set_paragraph_bold(paragraph, bold=True, font_size=12):
    if not paragraph.runs:
        run = paragraph.add_run("")
        run.font.bold = bold
        run.font.size = Pt(font_size)
        return

    for run in paragraph.runs:
        run.font.bold = bold
        run.font.size = Pt(font_size)

def _safe_str(value):
    if value is None:
        return ""
    if isinstance(value, int):
        return f"{value:,}"
    return str(value)

def _remove_paragraph(paragraph):
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def _cleanup_trailing_empty_paragraphs(doc: Document):
    for paragraph in reversed(doc.paragraphs):
        if paragraph.text.strip():
            break
        _remove_paragraph(paragraph)

def _set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for edge in ("top", "left", "bottom", "right"):
        element = tcBorders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tcBorders.append(element)

        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def _set_cell_font(cell, font_size=12):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(font_size)


def _set_cell_align_center(cell):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _replace_text_in_paragraph(paragraph, replacements: dict):
    full_text = "".join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text

    for key, value in replacements.items():
        token = f"{{{{{key}}}}}"
        if token in full_text:
            full_text = full_text.replace(token, _safe_str(value))

    if paragraph.runs:
        paragraph.runs[0].text = full_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = full_text


def _replace_text_in_cell(cell, replacements: dict):
    for paragraph in cell.paragraphs:
        _replace_text_in_paragraph(paragraph, replacements)


def _replace_text_in_table(table, replacements: dict):
    for row in table.rows:
        for cell in row.cells:
            _replace_text_in_cell(cell, replacements)


def _build_facility_payloads(project_data: dict):
    facilities = (
        project_data.get("prevention_facilities")
        or project_data.get("preventions")
        or []
    )
    payloads = []

    for idx, facility in enumerate(facilities, start=1):
        rows = []
        subtotal = 0

        if idx == 1:
            gw_item = facility.get("gw_item")
            vpn_item = facility.get("vpn_item")

            if gw_item:
                rows.append(gw_item)
                subtotal += int(gw_item.get("ITEM_AMOUNT", 0) or 0)

            if vpn_item:
                rows.append(vpn_item)
                subtotal += int(vpn_item.get("ITEM_AMOUNT", 0) or 0)

        for sensor in facility.get("sensors", []):
            rows.append(sensor)
            subtotal += int(sensor.get("ITEM_AMOUNT", 0) or 0)

        facility_no = facility.get("facility_no", f"방{idx}")
        raw_name = facility.get("prevention_name", f"{idx}번 방지시설")
        facility_capacity = facility.get("prevention_capacity", "")

        # raw_name 예: "방2 흡착에 의한 시설" 또는 "(방2) 흡착에 의한 시설"
        display_base = raw_name.strip()

        if display_base.startswith("("):
            name_only = display_base.split(")", 1)[-1].strip()
        else:
            name_only = display_base.split(" ", 1)[-1].strip() if " " in display_base else display_base

        if facility_capacity:
            display_name = f"({facility_no}) {name_only}({facility_capacity})"
        else:
            display_name = f"({facility_no}) {name_only}"

        payloads.append({
            "facility_no": facility_no,
            "facility_name": display_name,
            "rows": rows,
            "subtotal": subtotal,
        })

    return payloads


def _fill_summary_table(table, payloads: list):
    total_cost = sum(p["subtotal"] for p in payloads)
    payload_count = len(payloads)

    replacements = {
        "PREVENTION_COST_NO1": payloads[0]["subtotal"] if payload_count >= 1 else "",
        "PREVENTION_COST_NO2": payloads[1]["subtotal"] if payload_count >= 2 else "",
        "PREVENTION_COST_NO3": payloads[2]["subtotal"] if payload_count >= 3 else "",
        "TOTAL_COST": total_cost,
    }
    _replace_text_in_table(table, replacements)

    formula_text = "①"
    if payload_count == 2:
        formula_text = "①+②"
    elif payload_count >= 3:
        formula_text = "①+②+③"

    # 첫 번째 열 문구 수정
    for row in table.rows:
        if len(row.cells) >= 1 and "총 공사금액" in row.cells[0].text:
            row.cells[0].text = f"총 공사금액({formula_text})"
            _set_cell_font(row.cells[0], 12)
            _set_cell_bold(row.cells[0], True, 12)
            _set_cell_align_center(row.cells[0])

    # 요약표는 2행 구조(헤더/값) 기준
    if len(table.rows) >= 2:
        header_row = table.rows[0]
        value_row = table.rows[1]

        if payload_count == 1 and len(header_row.cells) >= 4 and len(value_row.cells) >= 4:
            # 헤더행: 2~4열 병합
            header_merged = header_row.cells[1].merge(header_row.cells[2])
            header_merged = header_merged.merge(header_row.cells[3])
            header_merged.text = "공사금액①"
            _set_cell_font(header_merged, 12)
            _set_cell_bold(header_merged, True, 12)
            _set_cell_align_center(header_merged)

            # 값행: 2~4열 병합
            value_merged = value_row.cells[1].merge(value_row.cells[2])
            value_merged = value_merged.merge(value_row.cells[3])
            value_merged.text = _safe_str(total_cost)
            _set_cell_font(value_merged, 12)
            _set_cell_bold(value_merged, False, 12)
            _set_cell_align_center(value_merged)

        elif payload_count == 2 and len(header_row.cells) >= 5 and len(value_row.cells) >= 5:
            header_row.cells[1].text = "공사금액①"
            header_row.cells[2].text = "공사금액②"

            # 남는 1열(3번 인덱스)을 비고 전 열과 병합
            header_merged = header_row.cells[2].merge(header_row.cells[3])
            header_merged.text = "공사금액②"

            value_merged = value_row.cells[2].merge(value_row.cells[3])
            value_merged.text = _safe_str(payloads[1]["subtotal"])

            _set_cell_font(header_row.cells[1], 12)
            _set_cell_bold(header_row.cells[1], True, 12)
            _set_cell_align_center(header_row.cells[1])

            _set_cell_font(header_merged, 12)
            _set_cell_bold(header_merged, True, 12)
            _set_cell_align_center(header_merged)

            _set_cell_font(value_row.cells[1], 12)
            _set_cell_bold(value_row.cells[1], False, 12)
            _set_cell_align_center(value_row.cells[1])

            _set_cell_font(value_merged, 12)
            _set_cell_bold(value_merged, False, 12)
            _set_cell_align_center(value_merged)

        elif payload_count >= 3 and len(header_row.cells) >= 4 and len(value_row.cells) >= 4:
            header_row.cells[1].text = "공사금액①"
            header_row.cells[2].text = "공사금액②"
            header_row.cells[3].text = "공사금액③"

            for idx in [1, 2, 3]:
                _set_cell_font(header_row.cells[idx], 12)
                _set_cell_bold(header_row.cells[idx], True, 12)
                _set_cell_align_center(header_row.cells[idx])

                _set_cell_font(value_row.cells[idx], 12)
                _set_cell_bold(value_row.cells[idx], False, 12)
                _set_cell_align_center(value_row.cells[idx])

    # 요약표 1행 전체 볼드
    if len(table.rows) >= 1:
        for cell in table.rows[0].cells:
            _set_cell_font(cell, 12)
            _set_cell_bold(cell, True, 12)
            _set_cell_align_center(cell)

    # 전체 셀 정리
    for row in table.rows:
        for cell in row.cells:
            cleaned = "\n".join(line.strip() for line in cell.text.splitlines() if line.strip())
            cell.text = cleaned
            _set_cell_font(cell, 12)
            _set_cell_align_center(cell)

    # 헤더행 볼드 재적용
    if len(table.rows) >= 1:
        for cell in table.rows[0].cells:
            _set_cell_bold(cell, True, 12)

def _remove_paragraph(paragraph):
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def _remove_unused_facility_sections(doc: Document, payload_count: int):
    remove_prefixes = []

    if payload_count == 1:
        remove_prefixes = ["2)", "3)"]
    elif payload_count == 2:
        remove_prefixes = ["3)"]

    if not remove_prefixes:
        return

    target_note = "사물인터넷 공사비용은 개별 단가를 적용하여 금액을 산출하여 정액지원함."
    paragraphs = list(doc.paragraphs)

    for i, paragraph in enumerate(paragraphs):
        text = " ".join(paragraph.text.replace("\xa0", " ").split()).strip()

        if any(text.startswith(prefix) and "사물인터넷(IoT) 설치내역" in text for prefix in remove_prefixes):
            # 제목 삭제
            _remove_paragraph(paragraph)

            # 바로 다음 안내문구 삭제
            if i + 1 < len(paragraphs):
                next_p = paragraphs[i + 1]
                next_text = " ".join(next_p.text.replace("\xa0", " ").split()).strip()
                if target_note in next_text:
                    _remove_paragraph(next_p)

            # 제목 앞 빈 문단 삭제
            if i - 1 >= 0:
                prev_p = paragraphs[i - 1]
                prev_text = prev_p.text.strip()
                if not prev_text:
                    _remove_paragraph(prev_p)

            # 안내문구 뒤 빈 문단 삭제
            if i + 2 < len(paragraphs):
                next2_p = paragraphs[i + 2]
                next2_text = next2_p.text.strip()
                if not next2_text:
                    _remove_paragraph(next2_p)

def _fill_section_heading_text(doc: Document, payloads: list):
    heading_map = {
        "PREVENTION_COST_NO1": payloads[0]["facility_name"] if len(payloads) >= 1 else "",
        "PREVENTION_COST_NO2": payloads[1]["facility_name"] if len(payloads) >= 2 else "",
        "PREVENTION_COST_NO3": payloads[2]["facility_name"] if len(payloads) >= 3 else "",
    }

    for paragraph in doc.paragraphs:
        _replace_text_in_paragraph(paragraph, heading_map)


def _fill_facility_table(table, facility_payload: dict | None):
    if not facility_payload:
        return

    rows_data = facility_payload["rows"]
    subtotal = facility_payload["subtotal"]

    template_row = None
    subtotal_row = None

    for row in table.rows:
        row_text = "".join(cell.text for cell in row.cells)

        if "{{ITEM_NAME}}" in row_text:
            template_row = row

        if "{{PREVENTION_SUBTOTAL}}" in row_text:
            subtotal_row = row

    if template_row is None or subtotal_row is None:
        return

    item_template_texts = [cell.text for cell in template_row.cells]
    subtotal_template_texts = [cell.text for cell in subtotal_row.cells]

    tbl = table._tbl
    tbl.remove(template_row._tr)
    tbl.remove(subtotal_row._tr)

    created_rows = []

    # 아이템 행 추가
    for item in rows_data:
        new_row = table.add_row()
        created_rows.append(new_row)

        for i, cell in enumerate(new_row.cells):
            if i >= len(item_template_texts):
                continue

            rendered = item_template_texts[i]
            rendered = rendered.replace("{{ITEM_NAME}}", _safe_str(item.get("ITEM_NAME")))
            rendered = rendered.replace("{{ITEM_UNIT_PRICE}}", _safe_str(item.get("ITEM_UNIT_PRICE")))
            rendered = rendered.replace("{{ITEM_QTY}}", _safe_str(item.get("ITEM_QTY")))
            rendered = rendered.replace("{{ITEM_AMOUNT}}", _safe_str(item.get("ITEM_AMOUNT")))

            # 구분열 / 비고열은 병합 예정
            if i == 0 or i == 5:
                cell.text = ""
            else:
                cell.text = rendered.strip()

            _set_cell_border(cell)
            _set_cell_font(cell, 12)
            _set_cell_align_center(cell)

    # 소계 행 추가
    subtotal_row_new = table.add_row()

    for i, cell in enumerate(subtotal_row_new.cells):
        if i >= len(subtotal_template_texts):
            continue

        rendered = subtotal_template_texts[i]
        rendered = rendered.replace("{{PREVENTION_SUBTOTAL}}", _safe_str(subtotal))

        cell.text = rendered
        _set_cell_border(cell)
        _set_cell_font(cell, 12)
        _set_cell_align_center(cell)

    # 소계행 단가~금액 병합 후 값 1개만 표시
    if len(subtotal_row_new.cells) >= 5:
        subtotal_row_new.cells[2].text = ""
        subtotal_row_new.cells[3].text = ""
        subtotal_row_new.cells[4].text = _safe_str(subtotal)

        merged_cell = subtotal_row_new.cells[2].merge(subtotal_row_new.cells[3])
        merged_cell = merged_cell.merge(subtotal_row_new.cells[4])
        merged_cell.text = _safe_str(subtotal)

        _set_cell_border(merged_cell)
        _set_cell_font(merged_cell, 12)
        _set_cell_align_center(merged_cell)

    # 좌/우 세로 병합
    if created_rows:
        first_row = created_rows[0]

        # 왼쪽 병합
        if len(first_row.cells) >= 1 and len(subtotal_row_new.cells) >= 1:
            merged_left = first_row.cells[0].merge(subtotal_row_new.cells[0])
            merged_left.text = "순\n공\n사\n원\n가"
            _set_cell_border(merged_left)
            _set_cell_font(merged_left, 12)
            _set_cell_align_center(merged_left)

        # 오른쪽 병합
        if len(first_row.cells) >= 6 and len(subtotal_row_new.cells) >= 6:
            merged_right = first_row.cells[5].merge(subtotal_row_new.cells[5])
            merged_right.text = "단위 : 원\n(VAT제외)"
            _set_cell_border(merged_right)
            _set_cell_font(merged_right, 12)
            _set_cell_align_center(merged_right)


def generate_doc_10040(template_path: Path, output_path: Path, project_data: dict):
    doc = Document(template_path)

    payloads = _build_facility_payloads(project_data)
    payload_count = len(payloads)

    if len(doc.tables) >= 1:
        _fill_summary_table(doc.tables[0], payloads)

    _fill_section_heading_text(doc, payloads)

    if len(doc.tables) >= 2 and payload_count >= 1:
        _fill_facility_table(doc.tables[1], payloads[0])

    if len(doc.tables) >= 3 and payload_count >= 2:
        _fill_facility_table(doc.tables[2], payloads[1])

    if len(doc.tables) >= 4 and payload_count >= 3:
        _fill_facility_table(doc.tables[3], payloads[2])

    # 사용하지 않는 세부표 삭제
    if payload_count == 1:
        while len(doc.tables) > 2:
            tbl = doc.tables[-1]._element
            tbl.getparent().remove(tbl)

    elif payload_count == 2:
        while len(doc.tables) > 3:
            tbl = doc.tables[-1]._element
            tbl.getparent().remove(tbl)

    _remove_unused_facility_sections(doc, payload_count)

    field_map = project_data.get("fields", {})
    for paragraph in doc.paragraphs:
        _replace_text_in_paragraph(paragraph, field_map)

    for table in doc.tables:
        _replace_text_in_table(table, field_map)

    # 안내문구 볼드 + 12pt 재적용
    for paragraph in doc.paragraphs:
        text = " ".join(paragraph.text.replace("\xa0", " ").split()).strip()

        if "방지시설 가동 여부를 확인하기 위하여 각각의 방지시설에 사물인터넷(IoT) 측정기기를 설치하여야 함." in text:
            _set_paragraph_bold(paragraph, True, 12)

        if "다만, IoT 게이트웨이, VPN 등은 중복 설치할 필요가 없으므로 동 기기는 1개의 설치비만 지원" in text:
            _set_paragraph_bold(paragraph, True, 12)

    _cleanup_trailing_empty_paragraphs(doc)
    doc.save(output_path)