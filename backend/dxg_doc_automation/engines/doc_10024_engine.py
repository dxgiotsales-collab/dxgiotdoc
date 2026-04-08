from __future__ import annotations

import copy
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.table import Table


PREVENTION_DETAIL_RULE = {
    "여과집진시설": {
        "title_text": "○ 여과집진시설 ○",
        "image_keys": ["temp_location", "dp_in_location", "dp_out_location"],
        "labels": ["온도계 설치 위치", "차압계 IN 설치 위치", "차압계 OUT 설치 위치"],
    },
    "흡착에 의한 시설": {
        "title_text": "○ 흡착에 의한 시설 ○",
        "image_keys": ["temp_location", "dp_in_location", "dp_out_location"],
        "labels": ["온도계 설치 위치", "차압계 IN 설치 위치", "차압계 OUT 설치 위치"],
    },
    "원심력 집진시설": {
        "title_text": "○ 원심력 집진시설 ○",
        "image_keys": [],
        "labels": [],
    },
    "세정집진시설": {
        "title_text": "○ 세정집진시설 ○",
        "image_keys": ["pump_ctrl_panel_out", "pump_ctrl_panel_in"],
        "labels": [
            "방지시설 펌프 설치위치\n(제어판넬 외함)",
            "방지시설 펌프 설치위치\n(제어판넬 내부)",
        ],
    },
    "전기집진시설": {
        "title_text": "○ 전기집진시설 ○",
        "image_keys": ["hv_ctrl_panel_out", "hv_ctrl_panel_in"],
        "labels": [
            "고전압발생기 전류계 설치위치\n(제어판넬 외함)",
            "고전압발생기 전류계 설치위치\n(제어판넬 내부)",
        ],
    },
    "여과 및 흡착에 의한 시설": {
        "title_text": "○ 여과 및 흡착에 의한 시설 ○",
        "image_keys": ["temp_location", "dp_in_location", "dp_out_location"],
        "labels": ["온도계 설치 위치", "차압계 IN 설치 위치", "차압계 OUT 설치 위치"],
    },
    "흡수에 의한 시설": {
        "title_text": "○ 흡수에 의한 시설 ○",
        "image_keys": ["ph_location", "pump_ctrl_panel_out", "pump_ctrl_panel_in"],
        "labels": [
            "방지시설 pH계 설치 위치",
            "방지시설 펌프 설치위치\n(제어판넬 외함)",
            "방지시설 펌프 설치위치\n(제어판넬 내부)",
        ],
    },
}


def _find_paragraph_index(doc: Document, token: str) -> Optional[int]:
    for i, p in enumerate(doc.paragraphs):
        if token in p.text:
            return i
    return None


def _remove_token_paragraph(doc: Document, token: str) -> None:
    idx = _find_paragraph_index(doc, token)
    if idx is None:
        return
    p_el = doc.paragraphs[idx]._element
    p_el.getparent().remove(p_el)


def _prevention_display_name(prevention: Dict[str, Any]) -> str:
    name = prevention.get("prevention_name", "") or prevention.get("prevention_method", "")
    cap = prevention.get("prevention_capacity", "")
    return f"{name}({cap})" if cap else name


def _emission_display_name(emission: Dict[str, Any]) -> str:
    name = emission.get("emission_name", "")
    cap = emission.get("emission_capacity", "")
    return f"{name}({cap})" if cap else name


def _clear_cell(cell) -> None:
    cell.text = ""
    if not cell.paragraphs:
        cell.add_paragraph("")


def _insert_image_into_cell(cell, image_path: Optional[str], width_cm: float = 4.8) -> None:
    _clear_cell(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if not image_path:
        return

    path = Path(image_path)

    # 🔥 절대경로 보정 (cwd 쓰지 말 것)
    if not path.exists():
        path = Path(__file__).resolve().parent.parent / path

    if path.exists():
        run = p.add_run()
        run.add_picture(str(path), width=Cm(width_cm))


def _set_cell_label(cell, text: str) -> None:
    _clear_cell(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(text or "")


def _collect_block_elements(doc: Document, start_token: str) -> List[Any]:
    body = doc._body._element
    children = list(body)
    start_idx = None

    for i, el in enumerate(children):
        if el.tag.endswith("}p"):
            text = "".join(t.text or "" for t in el.xpath(".//w:t"))
            if start_token in text:
                start_idx = i
                break

    if start_idx is None:
        raise ValueError(f"템플릿에서 시작 토큰을 찾지 못함: {start_token}")

    return children[start_idx + 1:]


def _remove_elements(doc: Document, elements: List[Any]) -> None:
    body = doc._body._element
    for el in elements:
        try:
            body.remove(el)
        except Exception:
            pass


def _append_elements_and_return(doc: Document, elements: List[Any]) -> List[Any]:
    body = doc._body._element
    appended = []
    for el in elements:
        new_el = copy.deepcopy(el)
        body.append(new_el)
        appended.append(new_el)
    return appended


def _get_first_table_from_elements(doc: Document, elements: List[Any]) -> Table:
    for el in elements:
        if el.tag.endswith("}tbl"):
            return Table(el, doc)
    raise ValueError("추가된 프로토타입 블록 안에서 바깥 wrapper 테이블을 찾지 못함.")


def _get_inner_tables_from_wrapper(wrapper_table: Table) -> List[Table]:
    outer_cell = wrapper_table.cell(0, 0)
    return outer_cell.tables


def _set_heading_in_wrapper(wrapper_table: Table, target_table: Table, text: str) -> None:
    outer_cell = wrapper_table.cell(0, 0)
    children = list(outer_cell._tc)
    target_el = target_table._element
    idx = children.index(target_el)

    for j in range(idx - 1, -1, -1):
        el = children[j]
        if el.tag.endswith("}p"):
            text_nodes = el.xpath(".//w:t")
            if text_nodes:
                text_nodes[0].text = text
                for extra in text_nodes[1:]:
                    extra.text = ""
            return


def _find_heading_before_table_in_wrapper(wrapper_table: Table, target_table: Table):
    outer_cell = wrapper_table.cell(0, 0)
    children = list(outer_cell._tc)
    target_el = target_table._element
    idx = children.index(target_el)

    for j in range(idx - 1, -1, -1):
        el = children[j]
        if el.tag.endswith("}p"):
            return el
    raise ValueError("표 위 제목 문단을 찾지 못함.")


def _clear_entire_table(table: Table) -> None:
    for r in range(len(table.rows)):
        for c in range(len(table.columns)):
            _clear_cell(table.cell(r, c))


def _remove_table_and_heading_in_wrapper(wrapper_table: Table, table: Table) -> None:
    outer_cell = wrapper_table.cell(0, 0)
    heading_el = _find_heading_before_table_in_wrapper(wrapper_table, table)
    table_el = table._element

    outer_cell._tc.remove(heading_el)
    outer_cell._tc.remove(table_el)


def _duplicate_table_block_in_wrapper(
    wrapper_table: Table,
    source_table: Table,
    insert_after_table: Table,
) -> Table:
    outer_cell = wrapper_table.cell(0, 0)

    source_heading_el = _find_heading_before_table_in_wrapper(wrapper_table, source_table)
    source_table_el = source_table._element

    new_heading_el = copy.deepcopy(source_heading_el)
    new_table_el = copy.deepcopy(source_table_el)

    insert_after_el = insert_after_table._element
    insert_after_el.addnext(new_heading_el)
    new_heading_el.addnext(new_table_el)

    return Table(new_table_el, outer_cell)


def _fill_common_table(table: Table, prevention: Dict[str, Any], image_width_cm: float) -> None:
    display_name = _prevention_display_name(prevention)
    common = prevention.get("common_images", {})

    _insert_image_into_cell(table.cell(0, 0), common.get("overview"), width_cm=image_width_cm)
    _insert_image_into_cell(table.cell(0, 1), common.get("gw_location"), width_cm=image_width_cm)
    _insert_image_into_cell(table.cell(0, 2), common.get("fan_ctrl_panel_out"), width_cm=image_width_cm)
    _insert_image_into_cell(table.cell(0, 3), common.get("fan_ctrl_panel_in"), width_cm=image_width_cm)

    _set_cell_label(table.cell(1, 0), display_name)
    _set_cell_label(table.cell(1, 1), f"{display_name}\nGATE WAY 설치 위치")
    _set_cell_label(table.cell(1, 2), "송풍시설(전류계) 설치 위치\n(제어판넬 외함)")
    _set_cell_label(table.cell(1, 3), "송풍시설(전류계) 설치 위치\n(제어판넬 내부)")


def _fill_detail_table_and_title(
    wrapper_table: Table,
    table: Table,
    prevention: Dict[str, Any],
    image_width_cm: float,
) -> None:

    method = prevention.get("prevention_method", "")

    if method == "원심력 집진시설":
        _set_heading_in_wrapper(wrapper_table, table, "○ 원심력 집진시설 ○")
        table._element.getparent().remove(table._element)
        return

    rule = PREVENTION_DETAIL_RULE.get(
        method,
        {
            "title_text": "○ 방지시설 상세 ○",
            "image_keys": [],
            "labels": [],
        },
    )

    _set_heading_in_wrapper(wrapper_table, table, rule["title_text"])

    image_keys = rule["image_keys"]
    labels = rule["labels"]
    detail = prevention.get("detail_images", {})
    cols = len(table.columns)

    for c in range(cols):
        _clear_cell(table.cell(0, c))
        _clear_cell(table.cell(1, c))

    for idx, key in enumerate(image_keys):
        if idx >= cols:
            break
        _insert_image_into_cell(table.cell(0, idx), detail.get(key), width_cm=image_width_cm)
        _set_cell_label(table.cell(1, idx), labels[idx])


def _fill_emission_table_in_wrapper(
    wrapper_table: Table,
    table: Table,
    prevention: Dict[str, Any],
    emission: Dict[str, Any],
    image_width_cm: float,
) -> None:
    _set_heading_in_wrapper(wrapper_table, table, "○ 배출시설 ○")

    display_name = _emission_display_name(emission)
    ctrl_out = emission.get("ctrl_out")
    ctrl_in = emission.get("ctrl_in")

    _insert_image_into_cell(table.cell(0, 0), emission.get("overview"), width_cm=image_width_cm)
    _insert_image_into_cell(table.cell(0, 1), ctrl_out, width_cm=image_width_cm)
    _insert_image_into_cell(table.cell(0, 2), ctrl_in, width_cm=image_width_cm)

    _set_cell_label(table.cell(1, 0), display_name)
    _set_cell_label(table.cell(1, 1), "배출시설(전류계) 설치 위치\n(제어판넬 외함)")
    _set_cell_label(table.cell(1, 2), "배출시설(전류계) 설치 위치\n(제어판넬 내부)")


def _fill_emission_table_2x4(
    wrapper_table: Table,
    table: Table,
    items: List[Dict[str, str]],
    image_width_cm: float,
) -> None:
    _set_heading_in_wrapper(wrapper_table, table, "○ 배출시설 ○")

    for c in range(4):
        _clear_cell(table.cell(0, c))
        _clear_cell(table.cell(1, c))

    for idx, item in enumerate(items[:4]):
        _insert_image_into_cell(table.cell(0, idx), item["img"], width_cm=image_width_cm)
        _set_cell_label(table.cell(1, idx), item["label"])


def _fill_emission_table_multi_using_common_template(
    wrapper_table: Table,
    common_table_template: Table,
    detail_table: Table,
    emission_table_template: Table,
    prevention: Dict[str, Any],
    emissions: List[Dict[str, Any]],
    image_width_cm: float,
) -> None:
    _remove_table_and_heading_in_wrapper(wrapper_table, emission_table_template)

    image_items: List[Dict[str, str]] = []

    for em in emissions:
        image_items.append(
            {
                "img": em.get("overview"),
                "label": _emission_display_name(em),
            }
        )

    image_items.append(
        {
            "img": prevention.get("emission_ctrl_panel_out"),
            "label": "배출시설(전류계) 설치 위치\n(제어판넬 외함)",
        }
    )
    image_items.append(
        {
            "img": prevention.get("emission_ctrl_panel_in"),
            "label": "배출시설(전류계) 설치 위치\n(제어판넬 내부)",
        }
    )

    chunks = [image_items[i:i + 4] for i in range(0, len(image_items), 4)]
    insert_after_table = detail_table

    for chunk in chunks:
        new_table = _duplicate_table_block_in_wrapper(
            wrapper_table=wrapper_table,
            source_table=common_table_template,
            insert_after_table=insert_after_table,
        )
        _fill_emission_table_2x4(
            wrapper_table=wrapper_table,
            table=new_table,
            items=chunk,
            image_width_cm=image_width_cm,
        )
        insert_after_table = new_table


def generate_doc_10024(
    template_path: str,
    output_dir: str,
    data: Dict[str, Any],
    *,
    image_width_cm: float = 4.8,
) -> str:
    template_path = str(template_path)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    start_token = "{{BLOCK_PREVENTION_PHOTO}}"
    doc = Document(template_path)

    current_block_elements = _collect_block_elements(doc, start_token)
    prototype_doc = Document(template_path)
    prototype_elements = _collect_block_elements(prototype_doc, start_token)

    _remove_token_paragraph(doc, start_token)
    _remove_elements(doc, current_block_elements)

    prevention_sections = data.get("prevention_sections", [])

    for idx, prevention in enumerate(prevention_sections):
        emissions = prevention.get("emissions", [])

        appended_elements = _append_elements_and_return(doc, prototype_elements)
        wrapper_table = _get_first_table_from_elements(doc, appended_elements)

        inner_tables = _get_inner_tables_from_wrapper(wrapper_table)
        if len(inner_tables) < 3:
            raise ValueError(f"중첩 표가 부족함. 필요>=3, 현재={len(inner_tables)}")

        common_table = inner_tables[0]
        detail_table = inner_tables[1]
        emission_table_template = inner_tables[2]

        _fill_common_table(common_table, prevention, image_width_cm=image_width_cm)
        _fill_detail_table_and_title(
            wrapper_table,
            detail_table,
            prevention,
            image_width_cm=image_width_cm,
        )

        if not emissions:
            _clear_entire_table(emission_table_template)
        else:
            _fill_emission_tables_repeat_1x3(
                wrapper_table=wrapper_table,
                emission_table_template=emission_table_template,
                detail_table=detail_table,
                prevention=prevention,
                emissions=emissions,
                image_width_cm=image_width_cm,
            )

        if idx < len(prevention_sections) - 1:
            doc.add_page_break()

    out_path = Path(output_dir) / "DOC_10024_output.docx"
    doc.save(str(out_path))
    return str(out_path)

def _fill_emission_tables_repeat_1x3(
    wrapper_table: Table,
    emission_table_template: Table,
    detail_table: Table,
    prevention: Dict[str, Any],
    emissions: List[Dict[str, Any]],
    image_width_cm: float,
) -> None:
    if not emissions:
        _clear_entire_table(emission_table_template)
        return

    _fill_emission_table_in_wrapper(
        wrapper_table,
        emission_table_template,
        prevention,
        emissions[0],
        image_width_cm=image_width_cm,
    )

    insert_after_table = emission_table_template

    for emission in emissions[1:]:
        new_table = _duplicate_table_block_in_wrapper(
            wrapper_table=wrapper_table,
            source_table=emission_table_template,
            insert_after_table=insert_after_table,
        )
        _fill_emission_table_in_wrapper(
            wrapper_table,
            new_table,
            prevention,
            emission,
            image_width_cm=image_width_cm,
        )
        insert_after_table = new_table