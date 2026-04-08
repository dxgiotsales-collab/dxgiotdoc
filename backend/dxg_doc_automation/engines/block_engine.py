from pathlib import Path
from docx import Document
from engines.token_engine import replace_token_in_paragraph, replace_token_in_table


def apply_block_rows_to_table(table, block_rows: list[dict]):
    if not block_rows:
        return

    template_row = table.rows[-1]
    tr = template_row._tr
    tbl = table._tbl
    tbl.remove(tr)

    for row_data in block_rows:
        new_row = table.add_row()
        for i, cell in enumerate(new_row.cells):
            if i >= len(template_row.cells):
                continue

            template_text = template_row.cells[i].text
            rendered_text = template_text

            for key, value in row_data.items():
                token = f"{{{{{key}}}}}"
                rendered_text = rendered_text.replace(token, str(value) if value is not None else "")

            cell.text = rendered_text


def build_doc_10040_rows(project_data: dict) -> list[dict]:
    rows = []
    prevention_facilities = project_data.get("prevention_facilities", [])

    for idx, facility in enumerate(prevention_facilities):
        sensors = facility.get("sensors", [])

        if idx == 0:
            gw = facility.get("gw_item")
            vpn = facility.get("vpn_item")
            if gw:
                rows.append(gw)
            if vpn:
                rows.append(vpn)

        for sensor in sensors:
            rows.append(sensor)

    return rows


def generate_block_doc(template_path: Path, output_path: Path, project_data: dict, doc_code: str):
    doc = Document(template_path)

    if doc_code == "DOC_10040":
        block_rows = build_doc_10040_rows(project_data)

        if doc.tables:
            apply_block_rows_to_table(doc.tables[0], block_rows)

    token_map = project_data.get("fields", {})
    for paragraph in doc.paragraphs:
        replace_token_in_paragraph(paragraph, token_map)

    for table in doc.tables:
        replace_token_in_table(table, token_map)

    doc.save(output_path)