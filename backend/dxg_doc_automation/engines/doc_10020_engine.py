from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.table import _Row
from docx.shared import Pt


def _safe_str(value):
    if value is None:
        return ""
    if isinstance(value, int):
        return f"{value:,}"
    return str(value)


def _set_cell_align_center(cell):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for run in paragraph.runs:
            run.font.size = Pt(11)


def _replace_text_in_paragraph(paragraph, replacements: dict):
    full_text = "".join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text

    for key, value in replacements.items():
        token = f"{{{{{key}}}}}"
        if token in full_text:
            full_text = full_text.replace(token, _safe_str(value))

    if paragraph.runs:
        paragraph.runs[0].text = full_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = full_text


def _replace_text_in_cell(cell, replacements: dict):
    for paragraph in cell.paragraphs:
        _replace_text_in_paragraph(paragraph, replacements)


def _replace_text_in_table(table, replacements: dict):
    for row in table.rows:
        for cell in row.cells:
            _replace_text_in_cell(cell, replacements)


def _fill_install_items_table(doc: Document, install_items: list[dict]):
    target_table = None
    template_row = None
    total_row = None

    for table in doc.tables:
        temp_template_row = None
        temp_total_row = None

        for row in table.rows:
            row_text = "".join(cell.text for cell in row.cells)

            if "{{ITEM_NAME}}" in row_text and "{{ITEM_UNIT_PRICE}}" in row_text:
                temp_template_row = row

            if "합계" in row_text and "{{TOTAL_COST}}" in row_text:
                temp_total_row = row

        if temp_template_row is not None and temp_total_row is not None:
            target_table = table
            template_row = temp_template_row
            total_row = temp_total_row
            break

    if target_table is None or template_row is None or total_row is None:
        return

    table = target_table
    tbl = table._tbl

    template_tr = deepcopy(template_row._tr)
    total_tr = deepcopy(total_row._tr)

    template_idx = list(tbl).index(template_row._tr)

    total_cost = sum(int(item.get("ITEM_AMOUNT", 0) or 0) for item in install_items)

    tbl.remove(template_row._tr)
    tbl.remove(total_row._tr)

    insert_idx = template_idx
    created_rows = []

    for item in install_items:
        new_tr = deepcopy(template_tr)
        tbl.insert(insert_idx, new_tr)

        new_row = _Row(new_tr, table)
        created_rows.append(new_row)

        replacements = {
            "ITEM_NAME": item.get("ITEM_NAME", ""),
            "ITEM_UNIT_PRICE": item.get("ITEM_UNIT_PRICE", ""),
            "ITEM_QTY": item.get("ITEM_QTY", ""),
            "ITEM_AMOUNT": item.get("ITEM_AMOUNT", ""),
        }

        for cell in new_row.cells:
            full_text = cell.text
            for key, value in replacements.items():
                full_text = full_text.replace(f"{{{{{key}}}}}", _safe_str(value))
            cell.text = full_text
            _set_cell_align_center(cell)

        insert_idx += 1

    new_total_tr = deepcopy(total_tr)
    tbl.insert(insert_idx, new_total_tr)
    new_total_row = _Row(new_total_tr, table)

    for cell in new_total_row.cells:
        full_text = cell.text.replace("{{TOTAL_COST}}", _safe_str(total_cost))
        cell.text = full_text
        _set_cell_align_center(cell)

    # 합계행 단가~금액(3칸) 병합
    if len(new_total_row.cells) >= 4:
        new_total_row.cells[1].text = ""
        new_total_row.cells[2].text = ""
        new_total_row.cells[3].text = _safe_str(total_cost)

        merged_cost = new_total_row.cells[1].merge(new_total_row.cells[2])
        merged_cost = merged_cost.merge(new_total_row.cells[3])
        merged_cost.text = _safe_str(total_cost)
        _set_cell_align_center(merged_cost)

    # 비고열 세로 병합
    if created_rows and len(created_rows[0].cells) >= 5 and len(new_total_row.cells) >= 5:
        remark_col_idx = len(new_total_row.cells) - 1

        merged_remark = created_rows[0].cells[remark_col_idx]
        for row in created_rows[1:]:
            merged_remark = merged_remark.merge(row.cells[remark_col_idx])

        merged_remark = merged_remark.merge(new_total_row.cells[remark_col_idx])
        merged_remark.text = "단위 : 원\n(VAT 제외)"
        _set_cell_align_center(merged_remark)


def generate_doc_10020(template_path: Path, output_path: Path, project_data: dict):
    doc = Document(template_path)

    field_map = project_data.get("fields", {})
    field_map_no_total = dict(field_map)
    field_map_no_total.pop("TOTAL_COST", None)

    install_items = project_data.get("install_items", [])

    for paragraph in doc.paragraphs:
        _replace_text_in_paragraph(paragraph, field_map_no_total)

    for table in doc.tables:
        _replace_text_in_table(table, field_map_no_total)

    _fill_install_items_table(doc, install_items)

    for paragraph in doc.paragraphs:
        _replace_text_in_paragraph(paragraph, field_map_no_total)

    for table in doc.tables:
        _replace_text_in_table(table, field_map_no_total)

    doc.save(output_path)