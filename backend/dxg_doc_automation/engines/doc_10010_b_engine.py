from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.table import _Row
from engines.doc_10010_a_engine import _fill_install_items_table


def _safe_str(value):
    if value is None:
        return ""
    return str(value)


def _set_cell_align_center(cell):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _replace_text_in_paragraph(paragraph, replacements: dict):
    full_text = "".join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text

    for key, value in replacements.items():
        token = f"{{{{{key}}}}}"
        if token in full_text:
            full_text = full_text.replace(token, _safe_str(value))

    if paragraph.runs:
        paragraph.runs[0].text = full_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = full_text


def _replace_text_in_cell(cell, replacements: dict):
    for paragraph in cell.paragraphs:
        _replace_text_in_paragraph(paragraph, replacements)


def _replace_text_in_table(table, replacements: dict):
    for row in table.rows:
        for cell in row.cells:
            _replace_text_in_cell(cell, replacements)

def _fill_site_facility_status_table_b(
    doc: Document,
    site_facility_status: list[dict],
    pollutants: list[dict],
    measurement_items: list[dict],
):

    target_table = None
    template_row = None
    header_row = None

    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            row_text = "".join(cell.text for cell in row.cells)
            if "{{BLOCK_SITE_FACILITY_STATUS}}" in row_text:
                target_table = table
                template_row = row
                if row_idx - 1 >= 0:
                    header_row = table.rows[row_idx - 1]
                break
        if target_table is not None:
            break

    if target_table is None or template_row is None or header_row is None:
        return

    table = target_table
    tbl = table._tbl

    header_texts = [c.text.strip().replace("\n", "").replace(" ", "") for c in header_row.cells]

    def find_col_idx(keyword: str):
        keyword = keyword.replace(" ", "")
        for i, txt in enumerate(header_texts):
            if keyword in txt:
                return i
        return None

    pollutant_type_col_idx = find_col_idx("대기오염물질종류")
    pollutant_amount_col_idx = find_col_idx("발생량")
    measure_col_idx = find_col_idx("최근자가측정결과") or find_col_idx("자가측정결과")

    def _format_measurement_item_line(item: dict) -> str:
        pollutant = _safe_str(item.get("pollutant", ""))
        amount = _safe_str(item.get("amount", ""))
        unit = _safe_str(item.get("unit", ""))
        date = _safe_str(item.get("date", ""))

        line1 = f"{pollutant} {amount}{unit}".strip()

        if date:
            try:
                yyyy, mm, dd = date.split("-")
                line2 = f"({yyyy[2:]}.{mm}.{dd})"
            except Exception:
                line2 = f"({date})"
            return f"{line1}\n{line2}".strip()

        return line1

    template_tr = deepcopy(template_row._tr)
    insert_idx = list(tbl).index(template_row._tr)
    tbl.remove(template_row._tr)

    created_rows = []
    row_count = max(len(site_facility_status), len(pollutants), len(measurement_items))

    for idx in range(row_count):
        facility = site_facility_status[idx] if idx < len(site_facility_status) else {}
        pollutant = pollutants[idx] if idx < len(pollutants) else {}
        measurement_item = measurement_items[idx] if idx < len(measurement_items) else {}

        new_tr = deepcopy(template_tr)
        tbl.insert(insert_idx, new_tr)
        new_row = _Row(new_tr, table)
        created_rows.append(new_row)

        replacements = {
            "BLOCK_SITE_FACILITY_STATUS": "",
            "EMISSION_FACILITY_NAME": facility.get("EMISSION_FACILITY_NAME", ""),
            "EMISSION_CAPACITY": facility.get("EMISSION_CAPACITY", ""),
            "EMISSION_QTY": facility.get("EMISSION_QTY", ""),
            "ITEM_POLLUTANT_TYPE": pollutant.get("ITEM_POLLUTANT_TYPE", ""),
            "ITEM_POLLUTANT_AMOUNT": pollutant.get("ITEM_POLLUTANT_AMOUNT", ""),
            "MEASURE_DATE": "",
        }

        for cell in new_row.cells:
            full_text = cell.text
            for key, value in replacements.items():
                full_text = full_text.replace(f"{{{{{key}}}}}", _safe_str(value))
            cell.text = full_text
            _set_cell_align_center(cell)

        insert_idx += 1

    if not created_rows:
        return

    # 👉 여기 넣기
    # 자가측정결과 열 병합 (measurement_items 기준)
    if measure_col_idx is not None and created_rows and measurement_items:
        total_rows = len(created_rows)
        total_items = len(measurement_items)

        base_span = total_rows // total_items
        extra = total_rows % total_items

        row_cursor = 0

        for item_idx, item in enumerate(measurement_items):
            span = base_span + (1 if item_idx < extra else 0)
            start_row = row_cursor
            end_row = min(total_rows - 1, row_cursor + span - 1)

            if start_row > end_row or start_row >= total_rows:
                continue

            merged_cell = created_rows[start_row].cells[measure_col_idx]

            for row_idx in range(start_row + 1, end_row + 1):
                created_rows[row_idx].cells[measure_col_idx].text = ""
                merged_cell = merged_cell.merge(created_rows[row_idx].cells[measure_col_idx])

            merged_cell.text = _format_measurement_item_line(item)
            _set_cell_align_center(merged_cell)

            row_cursor = end_row + 1


    # 오염물질 종류/발생량 병합 (pollutants 기준)
    if pollutant_type_col_idx is not None and pollutant_amount_col_idx is not None and pollutants:
        total_rows = len(created_rows)
        total_items = len(pollutants)

        base_span = total_rows // total_items
        extra = total_rows % total_items

        row_cursor = 0

        for item_idx, pollutant in enumerate(pollutants):
            span = base_span + (1 if item_idx < extra else 0)
            start_row = row_cursor
            end_row = min(total_rows - 1, row_cursor + span - 1)

            if start_row > end_row or start_row >= total_rows:
                continue

            pollutant_type = pollutant.get("ITEM_POLLUTANT_TYPE", "")
            pollutant_amount = pollutant.get("ITEM_POLLUTANT_AMOUNT", "")

            # 오염물질 종류 세로 병합
            merged_type_cell = created_rows[start_row].cells[pollutant_type_col_idx]
            for row_idx in range(start_row + 1, end_row + 1):
                created_rows[row_idx].cells[pollutant_type_col_idx].text = ""
                merged_type_cell = merged_type_cell.merge(created_rows[row_idx].cells[pollutant_type_col_idx])
            merged_type_cell.text = _safe_str(pollutant_type)
            _set_cell_align_center(merged_type_cell)

            # 발생량 세로 병합
            merged_amount_cell = created_rows[start_row].cells[pollutant_amount_col_idx]
            for row_idx in range(start_row + 1, end_row + 1):
                created_rows[row_idx].cells[pollutant_amount_col_idx].text = ""
                merged_amount_cell = merged_amount_cell.merge(created_rows[row_idx].cells[pollutant_amount_col_idx])
            merged_amount_cell.text = _safe_str(pollutant_amount)
            _set_cell_align_center(merged_amount_cell)

            row_cursor = end_row + 1


def generate_doc_10010_b(template_path: Path, output_path: Path, project_data: dict):
    doc = Document(template_path)

    field_map = project_data.get("fields", {})

    field_map_no_total = dict(field_map)
    field_map_no_total.pop("TOTAL_COST", None)
    field_map_no_total.pop("MEASURE_DATE", None)   # 추가

    pollutants = project_data.get("pollutants", [])

    if not pollutants:
        business = project_data.get("business", {}) or {}

        pollutant_type = business.get("pollutantType", "")
        pollutant_amount = business.get("pollutantAmount", "")

        if pollutant_type or pollutant_amount:
            type_list = [x.strip() for x in str(pollutant_type).split(",") if x.strip()]
            amount_list = [x.strip() for x in str(pollutant_amount).split(",") if x.strip()]

            row_count = max(len(type_list), len(amount_list))

            pollutants = [
                {
                    "ITEM_POLLUTANT_TYPE": type_list[i] if i < len(type_list) else "",
                    "ITEM_POLLUTANT_AMOUNT": amount_list[i] if i < len(amount_list) else "",
                }
                for i in range(row_count)
            ]

    business = project_data.get("business", {}) or {}
    measurement_items = business.get("measurementItems", []) or []

    site_facility_status = project_data.get("site_facility_status", [])
    install_items = project_data.get("install_items", [])

    # 1. 일반 토큰 치환 (TOTAL_COST, MEASURE_DATE 제외)
    for paragraph in doc.paragraphs:
        _replace_text_in_paragraph(paragraph, field_map_no_total)

    for table in doc.tables:
        _replace_text_in_table(table, field_map_no_total)

    # 2. 대기배출시설 현황
    _fill_site_facility_status_table_b(
        doc,
        site_facility_status,
        pollutants,
        measurement_items,
    )

    # 3. 설치대상 표
    _fill_install_items_table(doc, install_items)

    # 4. 마지막 일반 토큰 치환 (MEASURE_DATE 제외)
    final_field_map = dict(field_map)
    final_field_map.pop("MEASURE_DATE", None)

    for paragraph in doc.paragraphs:
        _replace_text_in_paragraph(paragraph, final_field_map)

    for table in doc.tables:
        _replace_text_in_table(table, final_field_map)

    doc.save(output_path)